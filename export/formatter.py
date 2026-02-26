"""
Модуль экспорта синопсиса в различные форматы: DOCX и PDF.
Поддерживает корпоративные шаблоны с логотипом и стандартными фразами.
"""

import io
import os
import re
from typing import Optional


def _strip_markdown(text: str) -> str:
    """Удаляет базовую Markdown-разметку для экспорта в plain-text форматы."""
    # Убираем заголовки
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Убираем жирный/курсив
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    # Убираем таблицы markdown (разделители)
    text = re.sub(r"^\|[-:| ]+\|$", "", text, flags=re.MULTILINE)
    return text


def export_docx(
    synopsis_md: str,
    title: str = "Синопсис протокола исследования биоэквивалентности",
    company_name: str = "Ifarma BE Study Planner",
    author: str = "Автоматически сгенерировано",
) -> bytes:
    """
    Генерирует DOCX-документ из Markdown-синопсиса.
    Поддерживает корпоративный шаблон с заголовком и метаданными.
    Возвращает байты DOCX-файла.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Стили ──────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Корпоративный заголовок ─────────────────────────────────────────
    hdr_para = doc.add_paragraph()
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hdr_para.add_run(company_name.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x84, 0xCC)

    # Разделитель
    doc.add_paragraph("─" * 60)

    # ── Заголовок документа ─────────────────────────────────────────────
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Метаданные
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{author}  |  Система: {company_name}").italic = True

    doc.add_paragraph("─" * 60)

    # ── Основной контент ─────────────────────────────────────────────────
    for line in synopsis_md.splitlines():
        stripped = line.rstrip()

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_formatted_run(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped.startswith("|"):
            # Строки таблицы — конвертируем в обычный текст
            if re.match(r"^\|[-:| ]+\|$", stripped):
                continue  # разделитель таблицы — пропускаем
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            p = doc.add_paragraph("\t".join(cells))
            p.style.font.size = Pt(10)
        elif stripped == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            _add_formatted_run(p, stripped)

    # ── Нижний колонтитул ────────────────────────────────────────────────
    doc.add_paragraph("─" * 60)
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run(
        "Документ сформирован автоматически системой Ifarma BE Study Planner. "
        "Требует верификации специалистом."
    ).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_formatted_run(paragraph, text: str) -> None:
    """Добавляет run с поддержкой **жирного** и *курсивного* текста."""
    # Парсим **bold** и *italic*
    parts = re.split(r"(\*{1,2}.+?\*{1,2})", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def export_pdf(
    synopsis_md: str,
    title: str = "Синопсис протокола исследования биоэквивалентности",
    company_name: str = "Ifarma BE Study Planner",
    author: str = "Автоматически сгенерировано",
) -> bytes:
    """
    Генерирует PDF-документ из Markdown-синопсиса.
    Использует fpdf2 с поддержкой Unicode (DejaVu шрифты при наличии).
    Возвращает байты PDF-файла.
    """
    from fpdf import FPDF

    # Ищем Unicode-шрифт в нескольких стандартных местах
    _FONT_CANDIDATES = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
        os.path.expanduser("~/Library/Fonts/DejaVuSans.ttf"),
    ]
    _FONT_BOLD_CANDIDATES = [p.replace("DejaVuSans.ttf", "DejaVuSans-Bold.ttf") for p in _FONT_CANDIDATES]
    _FONT_ITALIC_CANDIDATES = [p.replace("DejaVuSans.ttf", "DejaVuSans-Oblique.ttf") for p in _FONT_CANDIDATES]

    def _find_font(paths: list) -> Optional[str]:
        for p in paths:
            if os.path.exists(p):
                return p
        return None

    font_regular = _find_font(_FONT_CANDIDATES)
    font_bold    = _find_font(_FONT_BOLD_CANDIDATES)
    font_italic  = _find_font(_FONT_ITALIC_CANDIDATES)
    use_unicode  = bool(font_regular and font_bold and font_italic)
    _FONT_NAME   = "DejaVu" if use_unicode else "Helvetica"

    class CorpPDF(FPDF):
        def header(self):
            self.set_font(_FONT_NAME, "B", 9)
            self.set_text_color(0, 132, 204)
            self.cell(0, 8, company_name.upper(), align="C", new_x="LMARGIN", new_y="NEXT")
            self.set_draw_color(0, 132, 204)
            self.set_line_width(0.4)
            self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
            self.ln(2)

        def footer(self):
            self.set_y(-15)
            self.set_font(_FONT_NAME, "I", 8)
            self.set_text_color(120, 120, 120)
            self.cell(
                0, 10,
                "Dokument sozdan avtomaticheski. Trebuet verifikatsii spetsialistom.  "
                f"P. {self.page_no()}" if not use_unicode else
                "Документ сформирован автоматически. Требует верификации специалистом.  "
                f"Стр. {self.page_no()}",
                align="C",
            )

    pdf = CorpPDF()
    pdf.set_auto_page_break(auto=True, margin=18)

    if use_unicode:
        pdf.add_font(_FONT_NAME, "",  font_regular)
        pdf.add_font(_FONT_NAME, "B", font_bold)
        pdf.add_font(_FONT_NAME, "I", font_italic)

    pdf.add_page()

    # Заголовок документа
    _safe_title = title if use_unicode else _strip_markdown(title).encode("ascii", "replace").decode()
    pdf.set_font(_FONT_NAME, "B", 14)
    pdf.set_text_color(20, 20, 20)
    pdf.multi_cell(0, 8, _safe_title, align="C")
    pdf.ln(2)
    pdf.set_font(_FONT_NAME, "I", 9)
    pdf.set_text_color(100, 100, 100)
    _safe_author = author if use_unicode else author.encode("ascii", "replace").decode()
    pdf.multi_cell(0, 6, _safe_author, align="C")
    pdf.set_draw_color(180, 180, 180)
    pdf.set_line_width(0.3)
    pdf.line(pdf.l_margin, pdf.get_y() + 2, pdf.w - pdf.r_margin, pdf.get_y() + 2)
    pdf.ln(6)

    # Контент
    for line in synopsis_md.splitlines():
        stripped = line.rstrip()
        if not use_unicode:
            stripped = stripped.encode("ascii", "replace").decode()

        if stripped.startswith("# "):
            pdf.set_font(_FONT_NAME, "B", 13)
            pdf.set_text_color(0, 84, 160)
            pdf.multi_cell(0, 7, _strip_markdown(stripped), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
        elif stripped.startswith("## "):
            pdf.set_font(_FONT_NAME, "B", 11)
            pdf.set_text_color(0, 112, 192)
            pdf.multi_cell(0, 6, _strip_markdown(stripped), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
        elif stripped.startswith("### "):
            pdf.set_font(_FONT_NAME, "B", 10)
            pdf.set_text_color(50, 50, 50)
            pdf.multi_cell(0, 6, _strip_markdown(stripped), new_x="LMARGIN", new_y="NEXT")
        elif stripped.startswith("- ") or stripped.startswith("* "):
            pdf.set_font(_FONT_NAME, "", 10)
            pdf.set_text_color(30, 30, 30)
            content = _strip_markdown(stripped[2:])
            pdf.multi_cell(0, 5.5, f"  *  {content}", new_x="LMARGIN", new_y="NEXT")
        elif re.match(r"^\|[-:| ]+\|$", stripped):
            continue  # разделитель таблицы
        elif stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            pdf.set_font(_FONT_NAME, "", 9)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 5, "  |  ".join(cells), new_x="LMARGIN", new_y="NEXT")
        elif stripped == "":
            pdf.ln(3)
        else:
            pdf.set_font(_FONT_NAME, "", 10)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 5.5, _strip_markdown(stripped), new_x="LMARGIN", new_y="NEXT")

    return bytes(pdf.output())
